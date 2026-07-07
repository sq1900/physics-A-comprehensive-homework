#!/usr/bin/env python3
"""QA checks for Chinese course-report DOCX deliverables."""

from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


DEFAULT_PLACEHOLDERS = [
    "{{",
    "}}",
    "[[TOC]]",
    "待补",
    "待确认",
    "TODO",
    "FIXME",
]

DEFAULT_STALE_TERMS = [
    "实验一",
    "日志渲染",
]

DEFAULT_MOJIBAKE_TERMS = [
    "鐩綍",
    "瀹嬩綋",
    "榛戜綋",
    "鍥剧墖",
    "锛?",
]

COVER_MARKERS = ["课程报告", "实验题目", "学生姓名", "任课教师"]
SOURCE_PREFIX = "图片来源："
CITATION_PATTERN = re.compile(r"\[(?:\d+(?:\s*[-,，]\s*\d+)*)\]")
W_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}


def read_document_xml(docx_path: Path) -> str:
    with zipfile.ZipFile(docx_path) as zf:
        return zf.read("word/document.xml").decode("utf-8", errors="ignore")


def docx_text(doc: Document) -> str:
    parts: list[str] = []
    parts.extend(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def heading_counts(doc: Document) -> dict[str, int]:
    counts: dict[str, int] = {}
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name.startswith("Heading"):
            counts[style_name] = counts.get(style_name, 0) + 1
    return counts


def section_signature(doc: Document) -> list[dict[str, int]]:
    values = []
    for section in doc.sections:
        values.append(
            {
                "page_width": int(section.page_width),
                "page_height": int(section.page_height),
                "top_margin": int(section.top_margin),
                "bottom_margin": int(section.bottom_margin),
                "left_margin": int(section.left_margin),
                "right_margin": int(section.right_margin),
            }
        )
    return values


def style_names(doc: Document) -> set[str]:
    return {style.name for style in doc.styles}


def section_header_footer_text(doc: Document) -> str:
    parts: list[str] = []
    for section in doc.sections:
        for container in (section.header, section.footer):
            parts.extend(p.text for p in container.paragraphs if p.text.strip())
            for table in container.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells if cell.text.strip())
    return "\n".join(parts)


def find_terms(text: str, terms: list[str]) -> dict[str, int]:
    return {term: text.count(term) for term in terms if term and term in text}


def normalized_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def has_toc_field(xml: str) -> bool:
    return bool(re.search(r"TOC\s+\\o|TOC\\o|TOC ", xml))


def has_page_break_before_reference(xml: str) -> bool:
    reference_pos = xml.rfind("参考文献")
    if reference_pos < 0:
        return False
    before_reference = xml[max(0, reference_pos - 1500) : reference_pos]
    return bool(re.search(r'<w:br w:type="page"|<w:lastRenderedPageBreak', before_reference))


def has_body_page_number_restart(xml: str, doc: Document) -> bool:
    if not re.search(r'<w:pgNumType[^>]*w:start="1"', xml):
        return False

    try:
        root = ET.fromstring(xml)
    except ET.ParseError:
        return False

    body = root.find("w:body", W_NS)
    if body is None:
        return False

    body_paragraphs = [child for child in list(body) if child.tag.endswith("}p")]
    toc_seen = False
    section_break_after_toc = False

    for element, paragraph in zip(body_paragraphs, doc.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else ""
        has_toc_field = any("TOC" in (field.get(f"{{{W_NS['w']}}}instr") or "") for field in element.findall(".//w:fldSimple", W_NS))
        has_section_break = element.find("./w:pPr/w:sectPr", W_NS) is not None
        style_lower = style_name.lower()
        is_toc_content = has_toc_field or normalized_text(text) in {"目录", "目錄"} or style_lower.startswith("toc")
        is_body_heading = style_name.startswith("Heading") and text not in {"参考文献", "参考资料", "References"}

        if toc_seen and has_section_break:
            section_break_after_toc = True
        if toc_seen and is_body_heading:
            return section_break_after_toc
        if is_toc_content:
            toc_seen = True

    return False


def front_matter_has_page_fields(docx_path: Path) -> bool:
    try:
        with zipfile.ZipFile(docx_path) as zf:
            document_xml = zf.read("word/document.xml")
            rels_xml = zf.read("word/_rels/document.xml.rels")
            root = ET.fromstring(document_xml)
            rels_root = ET.fromstring(rels_xml)
            relmap = {rel.get("Id"): rel.get("Target") for rel in rels_root}
            sect_prs = root.findall(".//w:sectPr", W_NS)
            if len(sect_prs) <= 1:
                return False
            for sect_pr in sect_prs[:-1]:
                for footer_ref in sect_pr.findall("w:footerReference", W_NS):
                    rel_id = footer_ref.get(f"{{{W_NS['r']}}}id") if "r" in W_NS else None
                    if rel_id is None:
                        rel_id = footer_ref.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                    target = relmap.get(rel_id or "")
                    if not target:
                        continue
                    footer_name = target if target.startswith("word/") else f"word/{target}"
                    footer_xml = zf.read(footer_name).decode("utf-8", errors="ignore")
                    if "PAGE" in footer_xml:
                        return True
    except (KeyError, ET.ParseError, zipfile.BadZipFile):
        return False
    return False


def formal_figure_captions(text: str) -> list[str]:
    captions: list[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if re.match(r"^图\d+\.\d+\s+\S+", stripped):
            captions.append(stripped)
    return captions


def is_reference_list_paragraph(text: str, in_references: bool) -> bool:
    stripped = text.strip()
    if stripped in {"参考文献", "参考资料", "References"}:
        return True
    return in_references and bool(re.match(r"^\[\d+\]\s*\S+", stripped))


def plain_body_citation_markers(doc: Document) -> list[str]:
    plain: list[str] = []
    in_references = False
    for paragraph in doc.paragraphs:
        stripped = paragraph.text.strip()
        if stripped in {"参考文献", "参考资料", "References"}:
            in_references = True
            continue
        if is_reference_list_paragraph(stripped, in_references):
            continue
        for run in paragraph.runs:
            if run.font.superscript:
                continue
            plain.extend(match.group(0) for match in CITATION_PATTERN.finditer(run.text))
    return plain


def main() -> int:
    parser = argparse.ArgumentParser(description="QA a Chinese course-report DOCX.")
    parser.add_argument("--docx", type=Path, required=True)
    parser.add_argument("--require-toc", action="store_true")
    parser.add_argument("--require-cover", action="store_true")
    parser.add_argument("--min-images", type=int, default=0)
    parser.add_argument("--min-tables", type=int, default=0)
    parser.add_argument("--min-heading1", type=int, default=0)
    parser.add_argument("--require-reference-pagebreak", action="store_true")
    parser.add_argument("--require-body-page-start-1", action="store_true")
    parser.add_argument("--require-superscript-citations", action="store_true")
    parser.add_argument("--forbid-image-source-lines", action="store_true")
    parser.add_argument("--require-formal-figure-captions", action="store_true")
    parser.add_argument("--stale-term", action="append", default=[])
    parser.add_argument("--allow-placeholder", action="append", default=[])
    parser.add_argument("--template-fidelity-template", type=Path, help="Compare final DOCX against this user-provided template.")
    parser.add_argument("--template-marker", action="append", default=[], help="Template marker text expected to remain in final DOCX.")
    parser.add_argument("--allow-style-loss", action="store_true", help="Do not block when styles from the template are absent.")
    parser.add_argument("--allow-section-drift", action="store_true", help="Do not block when page setup differs from the template.")
    parser.add_argument("--json", action="store_true", help="Emit machine-readable JSON.")
    args = parser.parse_args()

    docx_path = args.docx.resolve()
    failures: list[str] = []
    warnings: list[str] = []

    if not docx_path.exists():
        print(f"BLOCK: DOCX not found: {docx_path}", file=sys.stderr)
        return 2

    doc = Document(str(docx_path))
    xml = read_document_xml(docx_path)
    text = docx_text(doc)
    headings = heading_counts(doc)
    inline_shapes = len(doc.inline_shapes)
    captions = formal_figure_captions(text)
    table_count = len(doc.tables)
    toc_field = has_toc_field(xml)
    reference_pagebreak = has_page_break_before_reference(xml)
    body_page_start_1 = has_body_page_number_restart(xml, doc)
    front_matter_page_fields = front_matter_has_page_fields(docx_path)
    plain_citations = plain_body_citation_markers(doc)
    template_fidelity: dict[str, object] = {}

    placeholder_terms = [t for t in DEFAULT_PLACEHOLDERS if t not in set(args.allow_placeholder)]
    placeholder_hits = find_terms(text, placeholder_terms)
    stale_hits = find_terms(text, DEFAULT_STALE_TERMS + args.stale_term)
    mojibake_hits = find_terms(text, DEFAULT_MOJIBAKE_TERMS)
    cover_marker_hits = find_terms(text, COVER_MARKERS)

    if placeholder_hits:
        failures.append(f"Placeholder residue found: {placeholder_hits}")
    if stale_hits:
        failures.append(f"Stale-topic terms found: {stale_hits}")
    if mojibake_hits:
        failures.append(f"Mojibake/encoding residue found: {mojibake_hits}")
    if args.require_toc and not toc_field:
        failures.append("Required automatic TOC field not found in DOCX XML.")
    if args.require_cover and not cover_marker_hits:
        failures.append("Required default/template cover markers not found.")
    if args.require_reference_pagebreak and not reference_pagebreak:
        failures.append("Required page break before 参考文献 not found.")
    if args.require_body_page_start_1 and not body_page_start_1:
        failures.append("Body page numbering does not restart at 1 after cover/TOC front matter.")
    if args.require_body_page_start_1 and front_matter_page_fields:
        failures.append("Cover/TOC front matter contains PAGE fields; the first visible page number must belong to the body.")
    if args.require_superscript_citations and plain_citations:
        failures.append(f"Plain body-sized citation markers found; expected superscript citations: {plain_citations}")
    if args.forbid_image_source_lines and SOURCE_PREFIX in text:
        failures.append("Image source lines are present in report body; keep provenance in sidecar files instead.")
    if args.require_formal_figure_captions and inline_shapes and len(captions) < inline_shapes:
        failures.append(
            f"Formal figure captions {len(captions)} < image count {inline_shapes}; expected captions like 图2.1 标题."
        )
    if inline_shapes < args.min_images:
        failures.append(f"Image count {inline_shapes} < required {args.min_images}.")
    if table_count < args.min_tables:
        failures.append(f"Table count {table_count} < required {args.min_tables}.")
    if headings.get("Heading 1", 0) < args.min_heading1:
        failures.append(f"Heading 1 count {headings.get('Heading 1', 0)} < required {args.min_heading1}.")
    if not headings:
        warnings.append("No Word heading styles detected; automatic TOC may not work.")

    for marker in args.template_marker:
        if marker and marker not in text:
            failures.append(f"Template marker not preserved in final DOCX: {marker}")

    if args.template_fidelity_template:
        template_doc = Document(str(args.template_fidelity_template.resolve()))
        template_sections = section_signature(template_doc)
        final_sections = section_signature(doc)
        template_styles = style_names(template_doc)
        final_styles = style_names(doc)
        missing_styles = sorted(template_styles - final_styles)
        template_hf_text = section_header_footer_text(template_doc)
        final_hf_text = section_header_footer_text(doc)
        template_fidelity = {
            "template": str(args.template_fidelity_template.resolve()),
            "section_match": template_sections == final_sections[: len(template_sections)],
            "missing_style_count": len(missing_styles),
            "missing_styles_sample": missing_styles[:20],
            "header_footer_preserved": (not template_hf_text.strip()) or template_hf_text.strip() in final_hf_text,
        }
        if not args.allow_section_drift and not template_fidelity["section_match"]:
            failures.append("Template page setup/section signature changed; pass --allow-section-drift only for approved fallback.")
        if not args.allow_style_loss and missing_styles:
            failures.append(f"Template styles missing from final DOCX: {missing_styles[:20]}")
        if not template_fidelity["header_footer_preserved"]:
            failures.append("Template header/footer text was not preserved in final DOCX.")

    result = {
        "docx": str(docx_path),
        "paragraphs": len(doc.paragraphs),
        "tables": table_count,
        "inline_shapes": inline_shapes,
        "formal_figure_captions": captions,
        "heading_counts": headings,
        "toc_field": toc_field,
        "reference_pagebreak": reference_pagebreak,
        "body_page_start_1": body_page_start_1,
        "front_matter_page_fields": front_matter_page_fields,
        "plain_body_citation_markers": plain_citations,
        "cover_marker_hits": cover_marker_hits,
        "placeholder_hits": placeholder_hits,
        "stale_hits": stale_hits,
        "mojibake_hits": mojibake_hits,
        "template_fidelity": template_fidelity,
        "warnings": warnings,
        "failures": failures,
        "status": "PASS" if not failures else "BLOCK",
    }

    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print(f"Status: {result['status']}")
        print(f"DOCX: {docx_path}")
        print(f"Paragraphs: {result['paragraphs']}")
        print(f"Tables: {table_count}")
        print(f"Images: {inline_shapes}")
        print(f"Headings: {headings}")
        print(f"TOC field: {toc_field}")
        print(f"Reference page break: {reference_pagebreak}")
        print(f"Body page start 1: {body_page_start_1}")
        print(f"Front matter PAGE fields: {front_matter_page_fields}")
        print(f"Cover markers: {cover_marker_hits}")
        for warning in warnings:
            print(f"WARN: {warning}")
        for failure in failures:
            print(f"BLOCK: {failure}")

    return 1 if failures else 0


if __name__ == "__main__":
    raise SystemExit(main())
