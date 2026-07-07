#!/usr/bin/env python3
"""Extract a DOCX Template Fidelity Contract for course-report templates."""

from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path

from docx import Document


def norm(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def para_records(doc: Document, limit: int) -> list[dict[str, str | int]]:
    records = []
    for index, paragraph in enumerate(doc.paragraphs[:limit]):
        records.append(
            {
                "index": index,
                "style": paragraph.style.name if paragraph.style is not None else "",
                "text": paragraph.text,
            }
        )
    return records


def section_records(doc: Document) -> list[dict[str, int]]:
    records = []
    for index, section in enumerate(doc.sections):
        records.append(
            {
                "index": index,
                "page_width": int(section.page_width),
                "page_height": int(section.page_height),
                "top_margin": int(section.top_margin),
                "bottom_margin": int(section.bottom_margin),
                "left_margin": int(section.left_margin),
                "right_margin": int(section.right_margin),
            }
        )
    return records


def table_records(doc: Document) -> list[dict[str, object]]:
    records = []
    for table_index, table in enumerate(doc.tables):
        rows = []
        for row in table.rows[:8]:
            rows.append([cell.text for cell in row.cells])
        records.append({"index": table_index, "rows": len(table.rows), "cols": len(table.columns), "sample": rows})
    return records


def media_count(path: Path) -> int:
    with zipfile.ZipFile(path) as zf:
        return sum(1 for name in zf.namelist() if name.startswith("word/media/"))


def header_footer_text(doc: Document) -> list[str]:
    values: list[str] = []
    for section in doc.sections:
        for container in (section.header, section.footer):
            values.extend(p.text for p in container.paragraphs if p.text.strip())
            for table in container.tables:
                for row in table.rows:
                    values.extend(cell.text for cell in row.cells if cell.text.strip())
    return values


def find_toc_and_body(paragraphs: list[dict[str, str | int]]) -> dict[str, int | None]:
    toc_index = None
    body_start_index = None
    for record in paragraphs:
        text = str(record["text"])
        compact = norm(text)
        if toc_index is None and compact in {"目录", "目錄"}:
            toc_index = int(record["index"])
            continue
        if toc_index is not None:
            if re.match(r"^第[一二三四五六七八九十]+章", compact) and not re.search(r"\d+$", compact):
                body_start_index = int(record["index"])
                break
            if re.match(r"^\d+(?:\.\d+){0,2}\S+", compact) and not re.search(r"\d+$", compact):
                body_start_index = int(record["index"])
                break
    return {"toc_index": toc_index, "body_start_index": body_start_index}


def build_contract(path: Path, paragraph_limit: int) -> dict[str, object]:
    doc = Document(str(path))
    paragraphs = para_records(doc, paragraph_limit)
    styles = sorted(style.name for style in doc.styles)
    return {
        "template": str(path.resolve()),
        "sections": section_records(doc),
        "styles": styles,
        "paragraph_sample": paragraphs,
        "tables": table_records(doc),
        "media_count": media_count(path),
        "headers_footers": header_footer_text(doc),
        "toc_body_detection": find_toc_and_body(paragraphs),
    }


def write_markdown(contract: dict[str, object], output: Path) -> None:
    lines = [
        "# Template Fidelity Contract",
        "",
        f"- Template: `{contract['template']}`",
        f"- Sections: {len(contract['sections'])}",
        f"- Styles: {len(contract['styles'])}",
        f"- Tables: {len(contract['tables'])}",
        f"- Media files: {contract['media_count']}",
        f"- TOC/body detection: `{contract['toc_body_detection']}`",
        "",
        "## Paragraph Sample",
        "",
        "| Index | Style | Text |",
        "|---:|---|---|",
    ]
    for record in contract["paragraph_sample"]:  # type: ignore[index]
        text = str(record["text"]).replace("|", "\\|")
        lines.append(f"| {record['index']} | {record['style']} | {text} |")
    lines.extend(["", "## Tables", ""])
    for table in contract["tables"]:  # type: ignore[index]
        lines.append(f"- Table {table['index']}: {table['rows']} rows x {table['cols']} cols")
    output.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description="Inspect a DOCX template and emit a Template Fidelity Contract.")
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--json-out", type=Path, required=True)
    parser.add_argument("--md-out", type=Path)
    parser.add_argument("--paragraph-limit", type=int, default=120)
    args = parser.parse_args()

    contract = build_contract(args.template, args.paragraph_limit)
    args.json_out.parent.mkdir(parents=True, exist_ok=True)
    args.json_out.write_text(json.dumps(contract, ensure_ascii=False, indent=2), encoding="utf-8")
    if args.md_out:
        args.md_out.parent.mkdir(parents=True, exist_ok=True)
        write_markdown(contract, args.md_out)
    print(args.json_out)
    if args.md_out:
        print(args.md_out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
