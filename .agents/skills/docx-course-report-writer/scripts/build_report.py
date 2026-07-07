#!/usr/bin/env python3
"""Build a Chinese course-report DOCX from a lightweight markdown draft."""

from __future__ import annotations

import argparse
from datetime import date
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SOURCE_PREFIX = "图片来源："
WIDTH_PREFIX = "图片宽度："
DEFAULT_TEMPLATE_PATH = Path(__file__).resolve().parents[1] / "skill-assets" / "default-course-report-template.docx"
DEFAULT_COVER_PARAGRAPHS = 11
TOC_RIGHT_TAB_CM = 15.35
CITATION_PATTERN = re.compile(r"\[(?:\d+(?:\s*[-,，]\s*\d+)*)\]")
CHAPTER_NUMERALS = {
    1: "一",
    2: "二",
    3: "三",
    4: "四",
    5: "五",
    6: "六",
    7: "七",
    8: "八",
    9: "九",
    10: "十",
    11: "十一",
    12: "十二",
    13: "十三",
    14: "十四",
    15: "十五",
}


def set_run_font(
    run,
    east_asia: str = "宋体",
    latin: str = "Times New Roman",
    size: float = 10.5,
    bold: bool | None = None,
    color: tuple[int, int, int] | None = None,
) -> None:
    run.font.name = latin
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def set_style_font(
    style,
    east_asia: str = "宋体",
    latin: str = "Times New Roman",
    size: float = 10.5,
    bold: bool | None = None,
) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def ensure_heading_styles(doc: Document) -> None:
    for level, size in ((1, 16), (2, 14), (3, 12)):
        name = f"Heading {level}"
        try:
            style = doc.styles[name]
        except KeyError:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles["Normal"]
        set_style_font(style, east_asia="黑体", size=size, bold=True)
        ppr = style._element.get_or_add_pPr()
        outline = ppr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            ppr.append(outline)
        outline.set(qn("w:val"), str(level - 1))


def get_or_create_paragraph_style(doc: Document, name: str):
    for style in doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH and style.name.lower() == name.lower():
            return style
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles["Normal"]
    return style


def ensure_toc_styles(doc: Document) -> None:
    specs = {
        "TOC 1": {"size": 11.0, "bold": True, "left": 0.0, "first": 0.0, "after": 4},
        "TOC 2": {"size": 10.5, "bold": False, "left": 0.55, "first": 0.0, "after": 3},
        "TOC 3": {"size": 10.0, "bold": False, "left": 1.1, "first": 0.0, "after": 2},
    }
    for name, spec in specs.items():
        style = get_or_create_paragraph_style(doc, name)
        set_style_font(style, east_asia="宋体", size=float(spec["size"]), bold=bool(spec["bold"]))
        fmt = style.paragraph_format
        fmt.left_indent = Cm(float(spec["left"]))
        fmt.first_line_indent = Cm(float(spec["first"]))
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(int(spec["after"]))
        fmt.line_spacing = 1.15
        fmt.tab_stops.clear_all()
        fmt.tab_stops.add_tab_stop(Cm(TOC_RIGHT_TAB_CM), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def configure_document(doc: Document, preserve_template_formatting: bool = False) -> None:
    """Configure generated report defaults.

    User-provided templates are authoritative layout artifacts. When preserving
    one, do not rewrite its page geometry or existing style definitions.
    """

    if preserve_template_formatting:
        for level in (1, 2, 3):
            name = f"Heading {level}"
            try:
                doc.styles[name]
            except KeyError:
                style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
                style.base_style = doc.styles["Normal"]
                set_style_font(style, east_asia="黑体", size={1: 16, 2: 14, 3: 12}[level], bold=True)
                ppr = style._element.get_or_add_pPr()
                outline = OxmlElement("w:outlineLvl")
                outline.set(qn("w:val"), str(level - 1))
                ppr.append(outline)
        return

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    ensure_heading_styles(doc)
    ensure_toc_styles(doc)


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    rows: list[str] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append(lines[i].strip())
        i += 1
    header = [c.strip() for c in rows[0].strip("|").split("|")]
    body = [[c.strip() for c in row.strip("|").split("|")] for row in rows[2:]]
    return header, body, i


def parse_width_cm(text: str) -> float:
    return float(text.replace("cm", "").strip())


def parse_tokens(text: str) -> list[tuple[str, Any]]:
    lines = text.splitlines()
    tokens: list[tuple[str, Any]] = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "{{REFERENCES}}":
            tokens.append(("references", None))
            i += 1
            continue

        if stripped == "{{PAGEBREAK}}":
            tokens.append(("pagebreak", None))
            i += 1
            continue

        if stripped.startswith("```"):
            language = stripped.strip("`").strip()
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            if i < len(lines):
                i += 1
            tokens.append(("code", {"language": language, "text": "\n".join(code_lines)}))
            continue

        if stripped.startswith("# "):
            tokens.append(("h1", stripped[2:].strip()))
            i += 1
            continue

        if stripped.startswith("## "):
            tokens.append(("h2", stripped[3:].strip()))
            i += 1
            continue

        if stripped.startswith("### "):
            tokens.append(("h3", stripped[4:].strip()))
            i += 1
            continue

        if stripped.startswith("!["):
            match = re.match(r"!\[(.+?)\]\((.+?)\)", stripped)
            if not match:
                raise ValueError(f"Invalid figure markdown: {stripped}")

            caption, path = match.group(1), match.group(2)
            source = ""
            width_cm = 13.8

            if i + 1 < len(lines) and lines[i + 1].strip().startswith(SOURCE_PREFIX):
                source = lines[i + 1].strip()
                i += 1

            if i + 1 < len(lines) and lines[i + 1].strip().startswith(WIDTH_PREFIX):
                width_line = lines[i + 1].strip().split("：", 1)[1]
                width_cm = parse_width_cm(width_line)
                i += 1

            tokens.append(("figure", {"caption": caption, "path": path, "source": source, "width_cm": width_cm}))
            i += 1
            continue

        if stripped.startswith("|"):
            header, body, i = parse_markdown_table(lines, i)
            tokens.append(("table", {"header": header, "rows": body}))
            continue

        block = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                i += 1
                break
            if nxt.startswith(("# ", "## ", "### ", "![", "|", "```")) or nxt in {"{{REFERENCES}}", "{{PAGEBREAK}}"}:
                break
            block.append(nxt)
            i += 1
        tokens.append(("p", " ".join(block)))

    return tokens


def preserve_template_opening(doc: Document, keep_paragraphs: int) -> None:
    body = doc._element.body
    seen = 0
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        if child.tag.endswith("}p"):
            seen += 1
            if seen <= keep_paragraphs:
                continue
        if element_has_section_properties(child):
            clear_element_text(child)
        else:
            body.remove(child)


def element_text(element) -> str:
    return "".join(node.text or "" for node in element.iter() if node.tag.endswith("}t"))


def element_has_section_properties(element) -> bool:
    return any(node.tag.endswith("}sectPr") for node in element.iter())


def clear_element_text(element) -> None:
    for node in element.iter():
        if node.tag.endswith("}t"):
            node.text = ""


def normalized_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def looks_like_toc_title(text: str) -> bool:
    return normalized_text(text) in {"目录", "目錄"}


def looks_like_body_start(text: str) -> bool:
    stripped = text.strip()
    compact = normalized_text(stripped)
    if re.match(r"^第[一二三四五六七八九十]+章", compact):
        return not re.search(r"\d+$", compact)
    if re.match(r"^\d+(?:\.\d+){0,2}\s*\S+", stripped):
        return not re.search(r"\s+\d+$", stripped)
    return False


def make_toc_field_paragraph(doc: Document):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run_el = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = "目录将在 Word 中更新。"
    run_el.append(text_el)
    fld.append(run_el)
    paragraph._p.append(fld)
    return paragraph._p


def set_section_page_number_start(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def reset_footer_paragraphs(footer):
    for child in list(footer._element):
        footer._element.remove(child)
    return footer.add_paragraph()


def clear_section_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    reset_footer_paragraphs(section.footer)


def add_page_number_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    paragraph = reset_footer_paragraphs(section.footer)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run_el = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = "1"
    run_el.append(text_el)
    fld.append(run_el)
    paragraph._p.append(fld)


def start_body_section_at_page_one(doc: Document) -> None:
    for section in doc.sections:
        clear_section_footer(section)
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_page_number_start(body_section, 1)
    add_page_number_footer(body_section)


def prepare_user_template_copy(doc: Document, fallback_opening_paragraphs: int, rebuild_toc: bool = True) -> dict[str, bool]:
    """Preserve a user template's opening while replacing stale body samples.

    The common report-template shape is cover -> TOC title/static TOC entries ->
    sample body headed by 第一章/1.1. This function keeps the template opening,
    removes old TOC entries and sample body, and inserts a real TOC field at the
    original TOC position when possible.
    """

    body = doc._element.body
    children = [child for child in list(body) if not child.tag.endswith("}sectPr")]
    toc_index: int | None = None
    body_start_index: int | None = None

    for index, child in enumerate(children):
        text = element_text(child)
        if toc_index is None and looks_like_toc_title(text):
            toc_index = index
            continue
        if toc_index is not None and looks_like_body_start(text):
            body_start_index = index
            break

    if body_start_index is None:
        preserve_template_opening(doc, fallback_opening_paragraphs)
        return {"auto_body_start": False, "toc_rebuilt": False, "preserved_opening": True}

    toc_rebuilt = False
    if rebuild_toc and toc_index is not None and toc_index < body_start_index:
        for child in children[toc_index + 1 : body_start_index]:
            if child.getparent() is not None:
                if element_has_section_properties(child):
                    clear_element_text(child)
                else:
                    body.remove(child)
        toc_field = make_toc_field_paragraph(doc)
        body.remove(toc_field)
        body.insert(toc_index + 1, toc_field)
        toc_rebuilt = True

    for child in children[body_start_index:]:
        if child.getparent() is not None:
            if element_has_section_properties(child):
                clear_element_text(child)
            else:
                body.remove(child)

    return {"auto_body_start": True, "toc_rebuilt": toc_rebuilt, "preserved_opening": True}


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if not child.tag.endswith("}sectPr"):
            body.remove(child)


def document_ends_with_section_break(doc: Document) -> bool:
    body = doc._element.body
    for child in reversed(list(body)):
        if child.tag.endswith("}sectPr"):
            continue
        return element_has_section_properties(child)
    return False


def resolve_template(args: argparse.Namespace) -> tuple[Path | None, bool]:
    if args.template:
        return args.template, False
    if not args.no_default_template and DEFAULT_TEMPLATE_PATH.exists():
        return DEFAULT_TEMPLATE_PATH, True
    return None, False


def parse_template_fields(raw_fields: list[str]) -> dict[str, str]:
    fields: dict[str, str] = {}
    for raw in raw_fields:
        if "=" not in raw:
            raise ValueError(f"--template-field must be KEY=VALUE, got: {raw}")
        key, value = raw.split("=", 1)
        fields[key.strip()] = value.strip()
    return fields


def template_value_map(args: argparse.Namespace, report_title: str) -> dict[str, str]:
    mapping = {
        "TITLE": report_title,
        "REPORT_TITLE": report_title,
        "实验题目": report_title,
        "COURSE": args.course.strip(),
        "课程名称": args.course.strip(),
        "COLLEGE": args.college.strip(),
        "学院系别": args.college.strip(),
        "MAJOR": args.major.strip(),
        "专业名称": args.major.strip(),
        "STUDENT_NAME": args.student_name.strip(),
        "学生姓名": args.student_name.strip(),
        "STUDENT_ID": args.student_id.strip(),
        "学生学号": args.student_id.strip(),
        "TEACHER": args.teacher.strip(),
        "任课教师": args.teacher.strip(),
        "DATE": args.date.strip() or current_date_text(),
        "完成日期": args.date.strip() or current_date_text(),
    }
    mapping.update(parse_template_fields(args.template_field))
    return {key: value for key, value in mapping.items() if value}


def replace_paragraph_text(paragraph, text: str) -> None:
    for run in list(paragraph.runs):
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def replace_text_placeholders(doc: Document, mapping: dict[str, str]) -> None:
    placeholder_variants = []
    for key, value in mapping.items():
        placeholder_variants.extend(
            [
                (f"{{{{{key}}}}}", value),
                (f"[[{key}]]", value),
                (f"【{key}】", value),
            ]
        )

    for paragraph in doc.paragraphs:
        text = paragraph.text
        for marker, value in placeholder_variants:
            if marker in text:
                text = text.replace(marker, value)
        compact = normalized_text(text)
        if compact in mapping and mapping[compact] and text.strip() == paragraph.text.strip():
            text = f"{paragraph.text}    {mapping[compact]}"
        if text != paragraph.text:
            replace_paragraph_text(paragraph, text)

    for table in doc.tables:
        for row in table.rows:
            for cell_index, cell in enumerate(row.cells):
                cell_text = normalized_text(cell.text)
                if cell_index + 1 < len(row.cells) and cell_text in mapping:
                    replace_paragraph_text(row.cells[cell_index + 1].paragraphs[0], mapping[cell_text])
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    for marker, value in placeholder_variants:
                        if marker in text:
                            text = text.replace(marker, value)
                    if text != paragraph.text:
                        replace_paragraph_text(paragraph, text)


def add_body_paragraph(doc: Document, text: str, indent: bool = True) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(6)
    if indent:
        paragraph.paragraph_format.first_line_indent = Pt(21)
    add_runs_with_citation_formatting(paragraph, text)


def add_runs_with_citation_formatting(paragraph, text: str) -> None:
    cursor = 0
    for match in CITATION_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=10.5)
        citation_run = paragraph.add_run(match.group(0))
        set_run_font(citation_run, size=9.0)
        citation_run.font.superscript = True
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=10.5)


def add_reference_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.05
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.left_indent = Cm(0.72)
    paragraph.paragraph_format.first_line_indent = Cm(-0.72)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.0)


def add_heading(doc: Document, text: str, level: int) -> None:
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    paragraph = doc.add_paragraph(style=style)
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt({1: 12, 2: 10, 3: 8}[level])
    paragraph.paragraph_format.space_after = Pt({1: 6, 2: 4, 3: 4}[level])
    run = paragraph.add_run(text)
    set_run_font(run, east_asia="黑体", size={1: 16, 2: 14, 3: 12}[level], bold=True)


def is_reference_heading(text: str) -> bool:
    normalized = re.sub(r"\s+", "", text)
    normalized = re.sub(r"^(第[一二三四五六七八九十]+章|\d+(?:\.\d+)*)", "", normalized)
    return normalized in {"参考文献", "参考资料", "References"}


class HeadingNumbering:
    def __init__(self) -> None:
        self.chapter = 0
        self.section = 0
        self.subsection = 0

    def format(self, text: str, level: int) -> str:
        if is_reference_heading(text):
            return "参考文献"

        stripped = re.sub(
            r"^(第[一二三四五六七八九十]+章\s*|\d+(?:\.\d+){0,2}\s*)",
            "",
            text.strip(),
        )

        if level == 1:
            self.chapter += 1
            self.section = 0
            self.subsection = 0
            chinese = CHAPTER_NUMERALS.get(self.chapter, str(self.chapter))
            return f"第{chinese}章 {stripped}"
        if level == 2:
            if self.chapter == 0:
                self.chapter = 1
            self.section += 1
            self.subsection = 0
            return f"{self.chapter}.{self.section} {stripped}"
        if level == 3:
            if self.chapter == 0:
                self.chapter = 1
            if self.section == 0:
                self.section = 1
            self.subsection += 1
            return f"{self.chapter}.{self.section}.{self.subsection} {stripped}"
        return stripped


def add_toc_field(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("目 录")
    set_run_font(run, east_asia="黑体", size=18, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run_el = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = "目录将在 Word 中更新。"
    run_el.append(text_el)
    fld.append(run_el)
    paragraph._p.append(fld)


def infer_report_title(draft_text: str, output: Path, explicit_title: str = "") -> str:
    if explicit_title.strip():
        return explicit_title.strip()
    for line in draft_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# ") and not is_reference_heading(stripped[2:].strip()):
            title = stripped[2:].strip()
            if not re.match(r"^(绪论|引言|实验名称|第一章|第[一二三四五六七八九十]+章)\b", title):
                return title
    return output.stem


def current_date_text() -> str:
    today = date.today()
    return f"{today.year} 年 {today.month} 月 {today.day} 日"


def populate_default_cover(doc: Document, args: argparse.Namespace, report_title: str) -> None:
    values = [
        "",
        "",
        f"《{args.course.strip() or '课程名称'}》",
        "课程报告",
        f"报告题目    {report_title}",
        f"学院系别    {args.college.strip() or '________________________________'}",
        f"专业名称    {args.major.strip() or '________________________________'}",
        f"学生姓名    {args.student_name.strip() or '________________________________'}",
        f"学生学号    {args.student_id.strip() or '________________________________'}",
        f"任课教师    {args.teacher.strip() or '________________________________'}",
        args.date.strip() or current_date_text(),
        "",
        "",
        "",
        "",
        "",
    ]

    for index, text in enumerate(values):
        if index >= len(doc.paragraphs):
            break
        paragraph = doc.paragraphs[index]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.first_line_indent = Pt(0)
        if index == 0:
            paragraph.paragraph_format.space_after = Pt(96)
        elif index == 2:
            paragraph.paragraph_format.space_after = Pt(16)
        elif index == 3:
            paragraph.paragraph_format.space_after = Pt(36)
        elif index == 4:
            paragraph.paragraph_format.space_after = Pt(34)
        elif 5 <= index <= 9:
            paragraph.paragraph_format.space_after = Pt(27)
        elif index == 10:
            paragraph.paragraph_format.space_before = Pt(48)
            paragraph.paragraph_format.space_after = Pt(0)
        elif 11 <= index <= 14:
            paragraph.paragraph_format.space_after = Pt(0)
        elif index == 15:
            paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        if index == 2:
            set_run_font(run, east_asia="黑体", size=18, bold=True)
        elif index == 3:
            set_run_font(run, east_asia="黑体", size=22, bold=True)
        elif index == 4:
            set_run_font(run, east_asia="宋体", size=13)
        elif 5 <= index <= 9:
            set_run_font(run, east_asia="宋体", size=12)
        elif index == 10:
            set_run_font(run, east_asia="宋体", size=12)
        else:
            set_run_font(run, east_asia="宋体", size=1 if not text else 12)


def formal_figure_caption(raw_caption: str, chapter: int, figure_index: int) -> str:
    caption = raw_caption.strip()
    caption = re.sub(r"^图\s*\d+(?:[.\-]\d+)?\s*", "", caption).strip()
    return f"图{chapter}.{figure_index} {caption}"


def add_figure(
    doc: Document,
    payload: dict[str, Any],
    root: Path,
    caption_text: str,
    include_source_lines: bool = False,
) -> None:
    fig_path = (root / str(payload["path"])).resolve()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.add_run().add_picture(str(fig_path), width=Cm(float(payload.get("width_cm", 13.8))))

    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.paragraph_format.space_after = Pt(6)
    run = caption_para.add_run(caption_text)
    set_run_font(run, size=10.5)

    source = str(payload.get("source", "")).strip()
    if source and include_source_lines:
        source_para = doc.add_paragraph()
        source_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        source_para.paragraph_format.space_after = Pt(6)
        run = source_para.add_run(source)
        set_run_font(run, size=9.5, color=(90, 90, 90))


def add_table(doc: Document, payload: dict[str, Any]) -> None:
    rows = payload["rows"]
    header = payload["header"]
    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for col, text in enumerate(header):
        cell = table.cell(0, col)
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "D9EAF7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=10.5, bold=True)

    for row_i, row in enumerate(rows, start=1):
        for col, text in enumerate(row):
            cell = table.cell(row_i, col)
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) <= 18 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(text))
            set_run_font(run, size=10.0)

    doc.add_paragraph()


def add_code_block(doc: Document, payload: dict[str, str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_shading(cell, "F5F5F5")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(payload["text"])
    set_run_font(run, east_asia="宋体", latin="Consolas", size=9.5)
    doc.add_paragraph()


def add_references(doc: Document, refs_path: Path, include_heading: bool = True) -> None:
    if not refs_path.exists():
        return
    if include_heading:
        doc.add_page_break()
        add_heading(doc, "参考文献", 1)
    for line in refs_path.read_text(encoding="utf-8").splitlines():
        if line.strip():
            add_reference_paragraph(doc, line.strip())


def build(args: argparse.Namespace) -> None:
    draft_text = args.draft.read_text(encoding="utf-8")
    report_title = infer_report_title(draft_text, args.output, args.title)
    template_path, used_default_template = resolve_template(args)
    user_template = bool(template_path and not used_default_template)
    user_template_status = {"auto_body_start": False, "toc_rebuilt": False, "preserved_opening": False}
    if template_path:
        doc = Document(str(template_path))
        if user_template:
            replace_text_placeholders(doc, template_value_map(args, report_title))
        if args.keep_template_body:
            pass
        elif user_template and args.drop_template_body:
            clear_document_body(doc)
        elif user_template:
            user_template_status = prepare_user_template_copy(
                doc,
                args.preserve_cover_paragraphs or args.user_template_opening_paragraphs,
                rebuild_toc=not args.no_toc,
            )
        elif args.preserve_cover_paragraphs > 0:
            preserve_template_opening(doc, args.preserve_cover_paragraphs)
        elif used_default_template and not args.drop_template_cover:
            preserve_template_opening(doc, DEFAULT_COVER_PARAGRAPHS)
            populate_default_cover(doc, args, report_title)
        else:
            clear_document_body(doc)
    else:
        doc = Document()

    configure_document(doc, preserve_template_formatting=user_template and not args.drop_template_body)

    if not args.no_toc:
        preserved_opening = template_path and (
            args.keep_template_body
            or (user_template and not args.drop_template_body)
            or args.preserve_cover_paragraphs > 0
            or (used_default_template and not args.drop_template_cover)
        )
        should_add_toc = not (user_template and user_template_status["toc_rebuilt"])
        if preserved_opening and should_add_toc and not document_ends_with_section_break(doc):
            doc.add_page_break()
        if should_add_toc:
            add_toc_field(doc)
        start_body_section_at_page_one(doc)

    tokens = parse_tokens(draft_text)
    numbering = HeadingNumbering()
    current_chapter = 1
    figure_counts: dict[int, int] = {}
    reference_heading_added = False
    for kind, payload in tokens:
        if kind == "h1":
            text = str(payload)
            if is_reference_heading(text):
                doc.add_page_break()
                add_heading(doc, "参考文献", 1)
                reference_heading_added = True
            else:
                heading_text = numbering.format(text, 1) if args.number_headings else text
                current_chapter = numbering.chapter if args.number_headings else max(current_chapter, 1)
                add_heading(doc, heading_text, 1)
                reference_heading_added = False
        elif kind == "h2":
            heading_text = numbering.format(str(payload), 2) if args.number_headings else str(payload)
            current_chapter = numbering.chapter if args.number_headings else max(current_chapter, 1)
            add_heading(doc, heading_text, 2)
            reference_heading_added = False
        elif kind == "h3":
            heading_text = numbering.format(str(payload), 3) if args.number_headings else str(payload)
            current_chapter = numbering.chapter if args.number_headings else max(current_chapter, 1)
            add_heading(doc, heading_text, 3)
            reference_heading_added = False
        elif kind == "p":
            add_body_paragraph(doc, str(payload))
            reference_heading_added = False
        elif kind == "figure":
            chapter = max(current_chapter, 1)
            figure_counts[chapter] = figure_counts.get(chapter, 0) + 1
            caption = (
                formal_figure_caption(str(payload["caption"]), chapter, figure_counts[chapter])
                if args.number_figures
                else str(payload["caption"])
            )
            add_figure(doc, payload, args.root, caption, include_source_lines=args.include_image_source_lines)
            reference_heading_added = False
        elif kind == "table":
            add_table(doc, payload)
            reference_heading_added = False
        elif kind == "code":
            add_code_block(doc, payload)
            reference_heading_added = False
        elif kind == "references":
            add_references(doc, args.refs, include_heading=not reference_heading_added)
            reference_heading_added = False
        elif kind == "pagebreak":
            doc.add_page_break()
            reference_heading_added = False

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a Chinese course-report DOCX.")
    parser.add_argument("--draft", type=Path, required=True)
    parser.add_argument("--refs", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--template", type=Path)
    parser.add_argument("--no-default-template", action="store_true")
    parser.add_argument("--keep-template-body", action="store_true")
    parser.add_argument(
        "--template-field",
        action="append",
        default=[],
        help="Extra user-template placeholder value as KEY=VALUE. Supports {{KEY}}, [[KEY]], 【KEY】 and table labels.",
    )
    parser.add_argument(
        "--drop-template-body",
        action="store_true",
        help="Explicitly discard a user-provided template body. Without this, --template uses copy-first preservation.",
    )
    parser.add_argument("--root", type=Path, default=Path("."))
    parser.add_argument("--preserve-cover-paragraphs", type=int, default=0)
    parser.add_argument(
        "--user-template-opening-paragraphs",
        type=int,
        default=DEFAULT_COVER_PARAGRAPHS,
        help="Leading paragraphs to preserve for user-provided templates when --drop-template-body is not set.",
    )
    parser.add_argument(
        "--drop-template-cover",
        action="store_true",
        help="When using the integrated default template, discard its visible cover and build only from styles/page setup.",
    )
    parser.add_argument("--no-toc", action="store_true")
    parser.add_argument("--title", default="", help="Report title used for the default template cover.")
    parser.add_argument("--course", default="", help="Course name used for the default template cover.")
    parser.add_argument("--college", default="", help="College/department used for the default template cover.")
    parser.add_argument("--major", default="", help="Major name used for the default template cover.")
    parser.add_argument("--student-name", default="", help="Student name used for the default template cover.")
    parser.add_argument("--student-id", default="", help="Student ID used for the default template cover.")
    parser.add_argument("--teacher", default="", help="Teacher name used for the default template cover.")
    parser.add_argument("--date", default="", help="Date text used for the default template cover.")
    parser.add_argument(
        "--no-heading-numbering",
        dest="number_headings",
        action="store_false",
        help="Disable automatic Chinese chapter and decimal section numbering.",
    )
    parser.add_argument(
        "--no-figure-numbering",
        dest="number_figures",
        action="store_false",
        help="Disable automatic formal figure captions such as 图2.1.",
    )
    parser.add_argument(
        "--include-image-source-lines",
        action="store_true",
        help="Render 图片来源 lines below captions. Default keeps provenance in sidecar attribution files only.",
    )
    parser.set_defaults(number_headings=True)
    parser.set_defaults(number_figures=True)
    build(parser.parse_args())


if __name__ == "__main__":
    main()
