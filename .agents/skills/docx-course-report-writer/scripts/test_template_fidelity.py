#!/usr/bin/env python3
"""Regression tests for user-provided template preservation."""

from __future__ import annotations

import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
BUILDER = ROOT / "scripts" / "build_report.py"
QA = ROOT / "scripts" / "qa_docx_report.py"


def all_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def make_template(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("课程报告")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "实验题目"
    table.cell(0, 1).text = "旧题目"
    table.cell(1, 0).text = "学生姓名"
    table.cell(1, 1).text = "{{STUDENT_NAME}}"
    doc.add_paragraph("目 录")
    doc.add_paragraph("第一章 1")
    doc.add_paragraph("第二章 2")
    doc.add_paragraph("第一章")
    doc.add_paragraph("旧样例正文，不应保留")
    doc.save(str(path))


def run_builder(*args: str) -> None:
    subprocess.run([sys.executable, str(BUILDER), *args], check=True)


def main() -> int:
    with tempfile.TemporaryDirectory() as td:
        work = Path(td)
        template = work / "template.docx"
        draft = work / "draft.md"
        refs = work / "refs.md"
        out_keep = work / "out_keep.docx"
        out_drop = work / "out_drop.docx"
        make_template(template)
        draft.write_text("# 实验任务\n\n正文内容。\n\n{{REFERENCES}}\n", encoding="utf-8")
        refs.write_text("[1] reference\n", encoding="utf-8")

        run_builder(
            "--draft",
            str(draft),
            "--refs",
            str(refs),
            "--output",
            str(out_keep),
            "--template",
            str(template),
            "--title",
            "新题目",
            "--student-name",
            "张三",
        )
        text = all_text(out_keep)
        assert "课程报告" in text
        assert "旧样例正文" not in text
        assert "正文内容" in text
        assert "新题目" in text
        assert "张三" in text
        assert text.count("目 录") == 1
        subprocess.run(
            [
                sys.executable,
                str(QA),
                "--docx",
                str(out_keep),
                "--template-fidelity-template",
                str(template),
                "--template-marker",
                "课程报告",
                "--allow-style-loss",
            ],
            check=True,
        )

        run_builder(
            "--draft",
            str(draft),
            "--refs",
            str(refs),
            "--output",
            str(out_drop),
            "--template",
            str(template),
            "--drop-template-body",
            "--no-toc",
        )
        dropped = all_text(out_drop)
        assert "课程报告" not in dropped
        assert "正文内容" in dropped

    print("PASS template fidelity regression")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
